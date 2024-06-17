import json
import boto3
import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from PIL import Image
from docx.oxml.ns import qn

# Import additional necessary classes or methods

DEBUG_MODE = False


def lambda_handler(event, context):
    """Lambda function handler for generating a CV document."""
    profile_picture_present = False
    resume_data = event["resume_data"]
    s3_bucket = "cvision-completed-resumes"
    s3_key = f'{event["upload_id"]}.docx'
    logo_path = "./Branding/Sirato_Logo_Color.png"
    image_path = "/tmp/profile.png"
    footer_text = (
        "Alle auf dem Dokument enthaltenen Informationen unterliegen den Allgemeinen Geschäftsbedingungen der "
        "Sirato Recruitment GmbH.\n\n"
        "Sirato Recruitment GmbH, Geschäftsführer: René Troche | Alter Hof 5, 80331 München | Deutschland\n"
        "Registergericht: Amtsgericht München | Registernummer: HRB 276100 | Ust-ID-Nr (gemäß § 27 a "
        "Umsatzsteuergesetz): DE353433498"
    )
    contact_details_present = event.get("contact_details_present", True)

    # Load the language file based on the event
    language_code = event.get("language", "en")
    global LANG_TEXT
    LANG_TEXT = load_translation(language_code)

    # Download image and logo if stored in S3
    if not DEBUG_MODE:
        s3 = boto3.client("s3")
        with open("/tmp/profile.png", "wb") as file:
            try:
                s3.download_fileobj(
                    "cv-profile-pictures", f'{event["upload_id"]}.png', file
                )
                profile_picture_present = True
            except s3.exceptions.NoSuchKey:
                print("No profile picture found with the specified key.")
            except s3.exceptions.ClientError as e:
                # A broader check for client issues, including access and connection errors
                if e.response["Error"]["Code"] == "404":
                    print("Profile picture not found.")
                else:
                    print(f"AWS Client error: {e}")
            except Exception as e:
                # Catching other unexpected exceptions
                print(f"An unexpected error occurred: {e}")

    # Create the sidebar sections
    sidebar_sections = [
        LanguageSection,
        CertificatesSection,
        SoftwareAndTechnologiesSection,
        SkillsAndCompetenciesSection,
        HobbiesSection,
    ]

    # Initialize CVBuilder with the loaded data and sidebar sections
    cv = CVBuilder(
        resume_data, sidebar_sections, contact_details_present, image_path, LANG_TEXT
    )

    BorderSection(cv, logo_path, footer_text)  # Header und Footer

    # Save the document to a temporary path
    if DEBUG_MODE:
        docx_path = "./tmp/resume.docx"
    else:
        docx_path = "/tmp/resume.docx"
    cv.save_document(docx_path)

    if DEBUG_MODE:
        print("DEBUG_MODE: ON")

    presigned_url = "Dummy"
    # Upload the document to S3
    if not DEBUG_MODE:
        s3.upload_file(docx_path, s3_bucket, s3_key)
        # Generate a presigned URL to download the resume
        presigned_url = s3.generate_presigned_url(
            "get_object", Params={"Bucket": s3_bucket, "Key": s3_key}, ExpiresIn=3600
        )
        print(presigned_url)

    return {
        "statusCode": 200,
        "body": json.dumps({"resume": presigned_url}),
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Credentials": True,
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type,X-Amz-Date,Authorization,X-Api-Key,X-Amz-Security-Token",
        },
    }


# Helper functions and classes

borders_visible = False


def disable_compatibility_mode(doc):
    """Disable compatibility mode in the document."""
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))

    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)

    compat_setting = OxmlElement("w:compatSetting")
    compat_setting.set(qn("w:name"), "compatibilityMode")
    compat_setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    compat_setting.set(qn("w:val"), "15")

    compat.append(compat_setting)


def adjust_page_layout(document):
    """Adjust the page layout of the document."""
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.0)
        section.right_margin = Inches(0.0)
        section.top_margin = Inches(0.0)
        section.bottom_margin = Inches(0.0)


def set_cell_background_color(cell, fill_color):
    """Set the background color of a cell."""
    hex_color = "{:02X}{:02X}{:02X}".format(*fill_color)
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), hex_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def apply_style_by_name(document, style_name):
    """Apply a style by its name to the document."""
    try:
        return document.styles[style_name]
    except KeyError:
        print(
            f"Style '{style_name}' not found in the document. Applying default style."
        )
        return document.styles["Normal"]  # Fallback to 'Normal' style if not found


def set_cell_borders(cell):
    """Remove the borders of a cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is not None:
        tc_pr.remove(borders)
    new_borders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")  # Set border value to 'nil' to remove the border
        new_borders.append(border)
    tc_pr.append(new_borders)


def set_row_height_to_page_height(table):
    """Set the row height to the height of the page."""
    row = table.rows[0]
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(
        qn("w:val"), str(int(Inches(11.69) * 1440))
    )  # 11.69 inches is the height of A4 paper
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def load_translation(lang_code):
    """Load the translation file based on the language code."""
    try:
        with open(f"Languages/{lang_code}.json", "r", encoding="utf-8") as file:
            print(f"Successfully loaded language file: {lang_code}.json")
            return json.load(file)
    except FileNotFoundError:
        print(
            f"Translation file not found: No such file or directory: {lang_code}.json"
        )
        return {}


def convert_to_compatible_jpeg(image_path):
    """Convert the image to a compatible JPEG format."""
    output_path = os.path.splitext(image_path)[0] + "_temp_picture.jpg"
    try:
        with Image.open(image_path) as img:
            img = img.convert("RGB")
            img.save(output_path, "JPEG", quality=85, optimize=True, progressive=True)
        return True, output_path
    except Exception as e:
        return False, f"Error converting image: {e}"


def validate_image(image_path):
    """Validate the image file."""
    try:
        with Image.open(image_path) as img:
            img.verify()
        return True, None
    except Exception as e:
        return False, f"Image validation failed for {image_path}: {e}"


def inches_to_emu(inches):
    """Convert inches to EMUs."""
    return int(inches * 914400)


class CVBuilder:
    """Base class for CVBuilder that defines general methods for the document."""

    def __init__(
        self,
        resume_data,
        sections,
        contact_details_present=True,
        image_path=None,
        lang_text={},
    ):
        self.doc = Document()
        disable_compatibility_mode(self.doc)
        adjust_page_layout(self.doc)
        self.resume_data = resume_data
        self.contact_details_present = contact_details_present
        self.sidebar_sections = sections
        self.image_path = image_path
        self.lang_text = lang_text
        self.add_sidebar(sections, image_path)

    def add_heading(self, text, size=15, style_name="Heading 1", cell=None):
        """Add a heading to the document."""
        heading_style = apply_style_by_name(self.doc, style_name)
        heading = cell.add_paragraph() if cell else self.doc.add_heading(level=1)
        run = heading.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x17, 0xC1, 0x61)  # Grünton
        heading.style = heading_style
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)

    def add_paragraph(
        self,
        text,
        size=10,
        bold=False,
        paragraph_space_after=6,
        space_after=True,
        space_after_pt=None,
        keep_with_next=False,
        style_name="Normal",
        cell=None,
    ):
        """Add a paragraph to the document."""
        paragraph_style = apply_style_by_name(self.doc, style_name)
        paragraph = cell.add_paragraph() if cell else self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold
        paragraph.paragraph_format.space_before = Pt(
            0
        )  # Ensure no space before paragraph
        if space_after:
            paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
        else:
            paragraph.paragraph_format.space_after = Pt(0)
        if space_after_pt is not None:
            paragraph.paragraph_format.space_after = Pt(space_after_pt)
        paragraph.paragraph_format.keep_with_next = keep_with_next
        paragraph.style = paragraph_style

    def add_paragraph_small(
        self,
        text,
        size=8,
        bold=False,
        paragraph_space_after=2,
        space_after=True,
        keep_with_next=False,
        cell=None,
    ):
        """Add a small paragraph to the document."""
        paragraph = cell.add_paragraph() if cell else self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(0x69, 0x69, 0x69)  # Dunkelgrauer Ton
        paragraph.paragraph_format.space_before = Pt(
            0
        )  # Ensure no space before paragraph
        if space_after:
            paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
        else:
            paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = keep_with_next

    def add_paragraph_custom(
        self,
        text,
        size=8,
        bold=False,
        italic=False,
        underline=False,
        color=None,
        paragraph_space_after=2,
        space_after=True,
        keep_with_next=False,
        cell=None,
    ):
        """Add a custom-styled paragraph to the document."""
        paragraph = self.doc.add_paragraph() if cell is None else cell.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.underline = underline

        if color:
            run.font.color.rgb = RGBColor(
                int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
            )

        paragraph.paragraph_format.space_before = Pt(
            0
        )  # Ensure no space before paragraph
        if space_after:
            paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
        else:
            paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.keep_with_next = keep_with_next

    def add_paragraph_to_cell(
        self,
        cell,
        text,
        size=11,
        bold=False,
        space_after=True,
        paragraph_space_after=None,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        keep_with_next=False,
        style_name="Normal",
    ):
        """Add a paragraph to a specific cell in the document."""
        paragraph_style = apply_style_by_name(self.doc, style_name)
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(
            0
        )  # Ensure no space before paragraph
        if space_after:
            paragraph.paragraph_format.space_after = Pt(6)
        else:
            paragraph.paragraph_format.space_after = Pt(0)
        if paragraph_space_after is not None:
            paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
        paragraph.paragraph_format.keep_with_next = keep_with_next
        paragraph.style = paragraph_style

    def add_bullet_points(
        self,
        bullet_points,
        size=10,
        bold=False,
        paragraph_space_after=2,
        space_after=True,
        keep_with_next=False,
        style_name="List Bullet",
        cell=None,
    ):
        """Add bullet points to the document."""
        bullet_style = apply_style_by_name(self.doc, style_name)
        for bullet_point in bullet_points:
            paragraph = (
                cell.add_paragraph(style=bullet_style)
                if cell
                else self.doc.add_paragraph(style=bullet_style)
            )
            run = paragraph.add_run(bullet_point)
            run.font.name = "Open Sans"
            run.font.size = Pt(size)
            run.bold = bold
            paragraph.paragraph_format.space_before = Pt(
                0
            )  # Ensure no space before paragraph
            if space_after:
                paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
            else:
                paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = keep_with_next

    def add_spacer_small(self, space_after_pt=2, cell=None):
        """Add a small spacer to the document."""
        if cell:
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(space_after_pt)
        elif self.doc.paragraphs:
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.paragraph_format.space_after = Pt(space_after_pt)

    def add_spacer_medium(self, space_after_pt=4, cell=None):
        """Add a medium spacer to the document."""
        if cell:
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(space_after_pt)
        elif self.doc.paragraphs:
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.paragraph_format.space_after = Pt(space_after_pt)

    def add_spacer_large(self, space_after_pt=6, cell=None):
        """Add a large spacer to the document."""
        if cell:
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(space_after_pt)
        elif self.doc.paragraphs:
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.paragraph_format.space_after = Pt(space_after_pt)

    def save_document(self, path):
        """Save the document to the specified path."""
        # Call the method to remove the last section before saving
        self.doc.save(path)
        print(f"Successfully created the document: {path}")

    def remove_trailing_paragraphs(self):
        """Remove trailing empty paragraphs that might cause an extra page."""
        while self.doc.paragraphs and self.doc.paragraphs[-1].text == "":
            p = self.doc.paragraphs[-1]
            p._element.getparent().remove(p)

    def remove_empty_paragraphs(self, element):
        """Remove any empty paragraphs from the given element (e.g., a cell or a document)."""
        for paragraph in element.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)

    def add_main_content(self, cell):
        """Add main content sections to the specified cell, ensuring no empty paragraphs before adding each section."""
        self.remove_empty_paragraphs(
            cell
        )  # Remove any existing empty paragraphs in the main content cell
        self.add_section_to_cell(ProfessionalSummarySection, cell)
        self.add_section_to_cell(WorkExperienceSection, cell)
        self.add_section_to_cell(EducationSection, cell)

    def add_sidebar_content(self, cell, sections, image_path):
        """Add sidebar content to the specified cell."""
        # Remove existing empty paragraphs in the sidebar cell
        self.remove_empty_paragraphs(cell)

        # Add profile picture at the top of the sidebar if present
        if os.path.exists(image_path):
            try:
                print(f"Image path: {image_path}")
                # Check if the image is valid
                valid, message = validate_image(image_path)
                if valid:
                    # Convert the image to a compatible JPEG format
                    converted, output_path = convert_to_compatible_jpeg(image_path)
                    if converted:
                        print(f"Adding picture from {output_path}")
                        # Add a paragraph for the image and center it
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(
                            output_path, width=Inches(2)
                        )  # Adjust the width as needed
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        print(f"Error converting image: {output_path}")
                else:
                    print(f"Invalid image: {message}")
            except Exception as e:
                print(f"Error adding profile picture: {e}")
        else:
            print(f"Image path does not exist: {image_path}")

        personal_info_section = PersonalInformationSection(
            self, self.contact_details_present
        )
        personal_info_section.add_to_cell(cell)

        for section in sections:
            section_instance = section(self)
            section_instance.add_to_cell(cell)
            self.add_spacer_to_cell(cell, space_after_pt=4)

    def add_sidebar(self, sections, image_path):
        """Create a table to layout the sidebar and main content."""
        table = self.doc.add_table(rows=1, cols=7)
        table.autofit = False

        # Define sidebar width and main content width based on the golden ratio
        total_width = 8.27  # Total width of A4 page in inches (approximately)
        ratio = 1 / 1.618
        sidebar_width = total_width * ratio / (1 + ratio)
        main_content_width = total_width - sidebar_width

        spacer_width = Inches(0.2)
        sidebar_actual_width = Inches(sidebar_width - 0.2 - 0.2)  # Subtract spacers
        main_actual_width = Inches(
            main_content_width
        )  # Subtract nothing to fit the size of the whole page

        # Set the column widths explicitly (leave as it is)
        table.columns[0].width = spacer_width  # Spacer 1
        table.columns[1].width = sidebar_actual_width  # Sidebar content | Spacer 2
        table.columns[2].width = spacer_width  # Spacer 3
        table.columns[3].width = spacer_width  # Spacer 4
        table.columns[4].width = main_actual_width  # Main content | Spacer 5
        table.columns[5].width = spacer_width  # Spacer 6

        for cell in table.columns[0].cells:
            cell.width = spacer_width
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[1].cells:
            cell.width = sidebar_actual_width
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[2].cells:
            cell.width = spacer_width
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[3].cells:
            cell.width = spacer_width
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[4].cells:
            cell.width = main_actual_width
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[5].cells:
            cell.width = spacer_width
            if borders_visible:
                set_cell_borders(cell)

        # Set vertical alignment for all cells
        for cell in table.rows[0].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Set the row height to fill the page using page height
        table.rows[0].height = Inches(11.69)  # Height of A4 page in inches

        # Left sidebar green spacer cell
        left_sidebar_cell = table.cell(0, 0)
        set_cell_background_color(left_sidebar_cell, (23, 193, 97))

        # Sidebar cell
        sidebar_cell = table.cell(0, 1)
        set_cell_background_color(sidebar_cell, (23, 193, 97))
        self.add_sidebar_content(sidebar_cell, sections, image_path)

        # Right sidebar green spacer cell
        right_sidebar_cell = table.cell(0, 2)
        set_cell_background_color(right_sidebar_cell, (23, 193, 97))

        # Main content cell
        main_content_cell = table.cell(0, 4)
        self.add_main_content(main_content_cell)

    def add_spacer_to_cell(self, cell, space_after_pt=12):
        """Add a spacer to the specified cell."""
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(space_after_pt)

    def add_section_to_cell(self, section_class, cell):
        """Add a section to the specified cell."""
        # Clear any existing empty paragraphs in the cell
        while cell.paragraphs and cell.paragraphs[0].text == "":
            p = cell.paragraphs[0]
            p._element.getparent().remove(p._element)

        section = section_class(self)
        section.add_to_cell(cell)

    def is_value_present(self, value):
        """Check if a value is present and valid."""
        if value in [None, "null", "NULL", "N/A", "", []]:
            return False
        return True


class BorderSection:
    """Class for adding header and footer sections to the document."""

    def __init__(
        self, cv_builder, logo_path, footer_text=None, footer_distance=Inches(0.2)
    ):
        self.cv_builder = cv_builder
        self.logo_path = logo_path
        self.footer_text = footer_text
        self.footer_distance = footer_distance
        self.add_header()
        if self.footer_text:
            self.add_footer()

    def add_header(self):
        """Add a header to the document."""
        section = self.cv_builder.doc.sections[0]
        header = section.header

        # Remove any default paragraph from the header
        for p in header.paragraphs:
            p.clear()

        # Set header distance and all margins to 0
        section.header_distance = Inches(0)
        section.top_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.bottom_margin = Inches(0)

        # Unlink header from previous sections
        header.is_linked_to_previous = False

        # Ensure header paragraph formatting is set to zero
        for para in header.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = Pt(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        total_width = Inches(8.27)  # Total width of A4 page in inches (approximately)

        # Adding table to header with specified width
        table = header.add_table(rows=1, cols=3, width=total_width)
        table.autofit = False

        # Define sidebar width and main content width based on the golden ratio
        ratio = 1 / 1.618
        sidebar_width = total_width * ratio / (1 + ratio)
        main_content_width = total_width - sidebar_width

        sidebar_actual_width = sidebar_width - Inches(0.4)  # Subtract spacers
        main_actual_width = main_content_width

        # Set the column widths explicitly in EMUs
        sidebar_width_emu = Emu(sidebar_actual_width)
        spacer_width_emu = Emu(Inches(0.4))
        main_content_width_emu = Emu(main_actual_width)

        table.columns[0].width = sidebar_width_emu  # Sidebar content
        table.columns[1].width = spacer_width_emu  # Spacer
        table.columns[2].width = main_content_width_emu  # Main content

        for cell in table.columns[0].cells:
            cell.width = sidebar_width_emu

        for cell in table.columns[1].cells:
            cell.width = spacer_width_emu

        for cell in table.columns[2].cells:
            cell.width = main_content_width_emu

        # Set vertical alignment for all cells
        for cell in table.rows[0].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Sidebar cell
        sidebar_cell = table.cell(0, 0)
        set_cell_background_color(sidebar_cell, (23, 193, 97))

        # Spacer cell
        spacer_cell = table.cell(0, 1)
        set_cell_background_color(spacer_cell, (23, 193, 97))

        # Add the logo to the header cell with vertical alignment adjustment
        header_content_cell = table.cell(0, 2)
        header_content_cell.vertical_alignment = (
            WD_ALIGN_VERTICAL.CENTER
        )  # Set vertical alignment to center
        paragraph = (
            header_content_cell.paragraphs[0]
            if header_content_cell.paragraphs
            else header_content_cell.add_paragraph()
        )

        # Adjust the spacing before the paragraph to lower the logo
        paragraph.paragraph_format.space_before = Pt(
            12
        )  # Adjust the value as needed for vertical positioning
        run = paragraph.add_run()
        run.add_picture(self.logo_path, width=Inches(1.5))  # Adjust the width as needed
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Remove any remaining empty paragraphs
        for p in header.paragraphs:
            if not p.text.strip() and len(p.runs) == 0:
                p._element.getparent().remove(p._element)

    def add_footer(self):
        """Add a footer to the document."""
        section = self.cv_builder.doc.sections[0]
        footer = section.footer

        # Remove default paragraphs
        for p in footer.paragraphs:
            p.clear()

        section.footer_distance = Inches(0)
        section.top_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.bottom_margin = Inches(0)

        footer.is_linked_to_previous = False

        for para in footer.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = Pt(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        total_width = 8.27  # Total width of A4 page in inches (approximately)
        ratio = 1 / 1.618
        sidebar_width = total_width * ratio / (1 + ratio)
        main_content_width = total_width - sidebar_width

        spacer_width = Inches(0.2)
        sidebar_actual_width = Inches(sidebar_width - 0.2 - 0.2)  # Subtract spacers
        main_actual_width = Inches(
            main_content_width
        )  # Subtract nothing to fit the size of the whole page

        # Adding table to footer with specified width
        table = footer.add_table(rows=1, cols=7, width=total_width)
        table.autofit = False

        # Set the column widths explicitly
        table.columns[0].width = Emu(spacer_width)
        table.columns[1].width = Emu(sidebar_actual_width)
        table.columns[2].width = Emu(spacer_width)
        table.columns[3].width = Emu(spacer_width)
        table.columns[4].width = Emu(main_actual_width)
        table.columns[5].width = Emu(spacer_width)
        table.columns[6].width = Emu(spacer_width)

        for cell in table.columns[0].cells:
            cell.width = Emu(spacer_width)
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[1].cells:
            cell.width = Emu(sidebar_actual_width)
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[2].cells:
            cell.width = Emu(spacer_width)
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[3].cells:
            cell.width = Emu(spacer_width)
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[4].cells:
            cell.width = Emu(main_actual_width)
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[5].cells:
            cell.width = Emu(spacer_width)
            if borders_visible:
                set_cell_borders(cell)

        for cell in table.columns[6].cells:
            cell.width = Emu(spacer_width)
            if borders_visible:
                set_cell_borders(cell)

        # Set vertical alignment for all cells
        for cell in table.rows[0].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Left sidebar green spacer cell
        left_sidebar_cell = table.cell(0, 0)
        set_cell_background_color(left_sidebar_cell, (23, 193, 97))

        # Sidebar cell
        sidebar_cell = table.cell(0, 1)
        set_cell_background_color(sidebar_cell, (23, 193, 97))

        # Right sidebar green spacer cell
        right_sidebar_cell = table.cell(0, 2)
        set_cell_background_color(right_sidebar_cell, (23, 193, 97))

        # Add footer content
        footer_content_cell = table.cell(0, 4)
        paragraph = (
            footer_content_cell.paragraphs[0]
            if footer_content_cell.paragraphs
            else footer_content_cell.add_paragraph()
        )
        run = paragraph.add_run()
        run.add_text("_" * 160)  # deprecated 175
        paragraph.add_run("\n" + self.footer_text)

        for run in paragraph.runs:
            run.font.size = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for p in footer.paragraphs:
            if not p.text.strip() and len(p.runs) == 0:
                p._element.getparent().remove(p._element)


class PersonalInformationSection:
    """Class for adding personal information section to the document."""

    def __init__(self, cv_builder, contact_details_present):
        self.cv_builder = cv_builder
        self.contact_details_present = contact_details_present

    def add(self):
        """Add personal information section to the document."""
        self.add_to_cell(self.cv_builder.doc.add_paragraph().add_run().element)

    def add_to_cell(self, cell):
        """Add personal information section to the specified cell."""
        personal_info = self.cv_builder.resume_data["Personal Information"]
        contact_info = self.cv_builder.resume_data["Contact Information"]

        # Add name
        full_name = ""
        if self.cv_builder.is_value_present(personal_info.get("Firstname")):
            full_name += personal_info.get("Firstname")
        if self.cv_builder.is_value_present(personal_info.get("Surname")):
            full_name += " " + personal_info.get("Surname")

        # Add the name to the cell and center-align it
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(full_name)
        run.font.name = "Open Sans"
        run.bold = True
        run.font.size = Pt(13)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)

        # Add an empty paragraph after the full name (better formatting)
        empty_paragraph = cell.add_paragraph()
        empty_paragraph.paragraph_format.space_after = Pt(0)

        # Add personal information
        info_present = any(
            [
                self.cv_builder.is_value_present(personal_info.get("Current Role")),
                self.cv_builder.is_value_present(personal_info.get("Birthday")),
                self.cv_builder.is_value_present(personal_info.get("Nationality")),
                self.cv_builder.is_value_present(personal_info.get("Marital Status")),
                self.cv_builder.is_value_present(personal_info.get("Availability")),
                self.cv_builder.is_value_present(
                    personal_info.get("Additional Information")
                ),
            ]
        )

        if info_present:
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(
                LANG_TEXT.get("PERSONAL_INFORMATION", "Personal Information")
            )
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)

            if self.cv_builder.is_value_present(personal_info.get("Birthday")):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(LANG_TEXT.get("BIRTHDAY", "Birthday") + ": ")
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(10)
                data_run = paragraph.add_run(personal_info.get("Birthday"))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)

            if self.cv_builder.is_value_present(personal_info.get("Nationality")):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(
                    LANG_TEXT.get("NATIONALITY", "Nationality") + ": "
                )
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(10)
                data_run = paragraph.add_run(personal_info.get("Nationality"))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)

            if self.cv_builder.is_value_present(personal_info.get("Marital Status")):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(
                    LANG_TEXT.get("MARITAL_STATUS", "Marital Status") + ": "
                )
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(10)
                data_run = paragraph.add_run(personal_info.get("Marital Status"))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)

            if self.cv_builder.is_value_present(personal_info.get("Availability")):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(
                    LANG_TEXT.get("AVAILABILITY", "Availability") + ": "
                )
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(10)
                data_run = paragraph.add_run(personal_info.get("Availability"))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)

            if self.cv_builder.is_value_present(
                personal_info.get("Additional Information")
            ):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(
                    LANG_TEXT.get("ADDITIONAL_INFORMATION", "Additional Information")
                    + ": "
                )
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(9)
                run.italic = True
                data_run = paragraph.add_run(
                    personal_info.get("Additional Information")
                )
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(9)
                data_run.italic = True
                paragraph.paragraph_format.space_after = Pt(2)

        # Add spacer between sections
        self.cv_builder.add_spacer_to_cell(cell, space_after_pt=10)

        # Add contact information
        if self.contact_details_present:
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(LANG_TEXT.get("CONTACT", "Contact"))
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)

            if self.cv_builder.is_value_present(contact_info.get("Address")):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(LANG_TEXT.get("ADDRESS", "Address") + ": ")
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(10)
                data_run = paragraph.add_run(contact_info.get("Address"))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)

            phone_numbers = [
                number
                for key, number in contact_info.items()
                if key.endswith("Phone Number")
                and self.cv_builder.is_value_present(number)
            ]
            if phone_numbers:
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(LANG_TEXT.get("PHONE", "Phone") + ": ")
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(10)
                data_run = paragraph.add_run(" | ".join(phone_numbers))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)

            if self.cv_builder.is_value_present(contact_info.get("Email")):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(LANG_TEXT.get("EMAIL", "Email") + ": ")
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(10)
                data_run = paragraph.add_run(contact_info.get("Email"))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(2)

            if self.cv_builder.is_value_present(
                contact_info.get("Additional Information")
            ):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(
                    LANG_TEXT.get("MORE_CONTACT_INFO", "Additional Contact Information")
                    + ": "
                )
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(9)
                run.italic = True
                data_run = paragraph.add_run(contact_info.get("Additional Information"))
                data_run.font.name = "Open Sans"
                data_run.font.size = Pt(9)
                data_run.italic = True
                paragraph.paragraph_format.space_after = Pt(2)

        # Add spacer at the end of the section
        self.cv_builder.add_spacer_to_cell(cell, space_after_pt=10)


class ProfessionalSummarySection:
    """Class for adding professional summary section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        """Add professional summary section to the document."""
        self.add_to_cell(self.cv_builder.doc.add_paragraph().add_run().element)

    def add_to_cell(self, cell):
        """Add professional summary section to the specified cell."""
        summary_data = self.cv_builder.resume_data.get("Professional Summary", {})
        summary_text = summary_data.get("Professional Summary Text", "")
        summary_bullet_points = summary_data.get(
            "Professional Summary Bullet Points", []
        )
        if self.cv_builder.is_value_present(summary_text):
            self.cv_builder.add_heading(
                LANG_TEXT.get("PROFESSIONAL_SUMMARY", "Berufliche Zusammenfassung"),
                size=13,
                cell=cell,
            )
            self.cv_builder.add_paragraph_to_cell(
                cell, summary_text, size=10, keep_with_next=True
            )
        if summary_bullet_points:
            self.cv_builder.add_bullet_points(
                summary_bullet_points, size=10, keep_with_next=False, cell=cell
            )
        # End Spacer
        self.cv_builder.add_spacer_large(space_after_pt=10, cell=cell)


class WorkExperienceSection:
    """Class for adding work experience section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        """Add work experience section to the document."""
        self.add_to_cell(self.cv_builder.doc.add_paragraph().add_run().element)

    def add_to_cell(self, cell):
        """Add work experience section to the specified cell."""
        self.cv_builder.add_heading(
            LANG_TEXT.get("WORK_EXPERIENCE", "Berufserfahrung"), size=13, cell=cell
        )
        experiences = self.cv_builder.resume_data.get("Working Experience", [])
        for item in experiences:
            title = item.get("Title", "NULL")
            company = item.get("Company", "")
            location = item.get("Location", "")
            start_date = item.get("Start Date", "")
            end_date = item.get("End Date", "heute")
            description = item.get("Description", "")
            bullet_points = item.get("Bullet Points", [])
            website = item.get("Website", "")
            additional_information = item.get("Additional Information", "")
            company_and_location_and_website = ""
            if self.cv_builder.is_value_present(company):
                company_and_location_and_website += item.get("Company", "")
            if self.cv_builder.is_value_present(location):
                if company_and_location_and_website:
                    company_and_location_and_website += " | "
                company_and_location_and_website += item.get("Location", "")
            if self.cv_builder.is_value_present(website):
                if company_and_location_and_website:
                    company_and_location_and_website += " | "
                company_and_location_and_website += item.get("Website", "")
            date_range = start_date
            if self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range += f" - {end_date}"
            elif not self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range = end_date
            if self.cv_builder.is_value_present(date_range):
                self.cv_builder.add_paragraph_small(
                    date_range,
                    size=8,
                    space_after=False,
                    keep_with_next=True,
                    cell=cell,
                )
            if self.cv_builder.is_value_present(title):
                self.cv_builder.add_paragraph_to_cell(
                    cell,
                    title,
                    size=11,
                    bold=True,
                    space_after=False,
                    keep_with_next=True,
                )
            if self.cv_builder.is_value_present(company_and_location_and_website):
                self.cv_builder.add_paragraph_small(
                    company_and_location_and_website,
                    space_after=True,
                    bold=True,
                    keep_with_next=True,
                    cell=cell,
                )
            if self.cv_builder.is_value_present(description):
                self.cv_builder.add_paragraph_to_cell(
                    cell, description, size=10, space_after=False, keep_with_next=True
                )
            if bullet_points:
                self.cv_builder.add_bullet_points(
                    bullet_points, keep_with_next=False, cell=cell
                )
            if self.cv_builder.is_value_present(additional_information):
                self.cv_builder.add_paragraph_custom(
                    additional_information,
                    size=9,
                    italic=True,
                    space_after=False,
                    keep_with_next=False,
                    cell=cell,
                )
            self.cv_builder.add_spacer_large(space_after_pt=12, cell=cell)


class EducationSection:
    """Class for adding education section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        """Add education section to the document."""
        self.add_to_cell(self.cv_builder.doc.add_paragraph().add_run().element)

    def add_to_cell(self, cell):
        """Add education section to the specified cell."""
        self.cv_builder.add_heading(
            LANG_TEXT.get("EDUCATION", "Bildung"), size=13, cell=cell
        )
        education_items = self.cv_builder.resume_data.get("Education", [])
        for item in education_items:
            diploma = item.get("Diploma", "NULL")
            institution = item.get("Institution", "")
            start_date = item.get("Start Date", "")
            end_date = item.get("End Date", "heute")
            grade = item.get("Grade", "")
            location = item.get("Location", "")
            website = item.get("Website", "")
            description = item.get("Description", "")
            bullet_points = item.get("Bullet Points", [])
            additional_information = item.get("Additional Information")
            institution_and_location_and_website = ""
            if self.cv_builder.is_value_present(institution):
                institution_and_location_and_website += item.get("Institution", "")
            if self.cv_builder.is_value_present(location):
                if institution_and_location_and_website:
                    institution_and_location_and_website += " | "
                institution_and_location_and_website += item.get("Location", "")
            if self.cv_builder.is_value_present(website):
                if institution_and_location_and_website:
                    institution_and_location_and_website += " | "
                institution_and_location_and_website += item.get("Website", "")
            date_range = start_date
            if self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range += f" - {end_date}"
            elif not self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range = end_date
            if self.cv_builder.is_value_present(date_range):
                self.cv_builder.add_paragraph_small(
                    date_range,
                    size=8,
                    space_after=False,
                    keep_with_next=True,
                    cell=cell,
                )
            if self.cv_builder.is_value_present(diploma):
                self.cv_builder.add_paragraph_to_cell(
                    cell,
                    diploma,
                    size=11,
                    bold=True,
                    space_after=False,
                    keep_with_next=True,
                )
            if self.cv_builder.is_value_present(institution_and_location_and_website):
                self.cv_builder.add_paragraph_small(
                    institution_and_location_and_website,
                    space_after=False,
                    bold=True,
                    keep_with_next=True,
                    cell=cell,
                )
            if self.cv_builder.is_value_present(grade):
                self.cv_builder.add_paragraph_to_cell(
                    cell,
                    LANG_TEXT.get("GRADE", "Note") + " " + grade,
                    size=10,
                    space_after=False,
                    keep_with_next=True,
                )
            if self.cv_builder.is_value_present(description):
                self.cv_builder.add_paragraph_to_cell(
                    cell, description, size=10, space_after=False, keep_with_next=True
                )
            if bullet_points:
                self.cv_builder.add_bullet_points(
                    bullet_points, keep_with_next=False, cell=cell
                )
            if self.cv_builder.is_value_present(additional_information):
                self.cv_builder.add_paragraph_custom(
                    additional_information,
                    size=9,
                    italic=True,
                    space_after=False,
                    keep_with_next=False,
                    cell=cell,
                )
            # In-Between spacer
            self.cv_builder.add_spacer_large(space_after_pt=8, cell=cell)
        # End Spacer
        self.cv_builder.add_spacer_large(space_after_pt=10, cell=cell)


class LanguageSection:
    """Class for adding language section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        """Add language section to the document."""
        languages = self.cv_builder.resume_data.get("Languages", [])
        valid_languages = []
        for language in languages:
            if self.cv_builder.is_value_present(language.get("Name")):
                name = language["Name"]
                level = (
                    language["Level"]
                    if self.cv_builder.is_value_present(language.get("Level"))
                    else ""
                )
                valid_languages.append((name, level))
        if valid_languages:
            self.cv_builder.add_heading(LANG_TEXT.get("LANGUAGES", "Sprachen"), size=12)
            for name, level in valid_languages:
                self.add_language_paragraph(name, level)

    def add_language_paragraph(self, name, level):
        """Add a language paragraph to the document."""
        paragraph = self.cv_builder.doc.add_paragraph()
        if level:
            run = paragraph.add_run(f"{name}: ")
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(11)
            run = paragraph.add_run(level)
            run.font.name = "Open Sans"
            run.bold = False
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(name)
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(11)
        paragraph.paragraph_format.space_after = Pt(4)

    def add_to_cell(self, cell):
        """Add language section to the specified cell."""
        languages = self.cv_builder.resume_data.get("Languages", [])
        valid_languages = []
        for language in languages:
            if self.cv_builder.is_value_present(language.get("Name")):
                name = language["Name"]
                level = (
                    language["Level"]
                    if self.cv_builder.is_value_present(language.get("Level"))
                    else ""
                )
                valid_languages.append((name, level))
        if valid_languages:
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(LANG_TEXT.get("LANGUAGES", "Sprachen"))
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            for name, level in valid_languages:
                self.add_language_paragraph_to_cell(cell, name, level)
        #  self.cv_builder.add_spacer_to_cell(cell, space_after_pt=10)  # Adjusted space after the section

    def add_language_paragraph_to_cell(self, cell, name, level):
        """Add a language paragraph to the specified cell."""
        paragraph = cell.add_paragraph()
        if level:
            run = paragraph.add_run(f"{name}: ")
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(11)
            run = paragraph.add_run(level)
            run.font.name = "Open Sans"
            run.bold = False
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(name)
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(11)
        paragraph.paragraph_format.space_after = Pt(4)


class CertificatesSection:
    """Class for adding certificates section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self, cell):
        """Add certificates section to the document."""
        certificates = self.cv_builder.resume_data.get("Certificates", [])
        if certificates:
            self.cv_builder.add_heading(
                LANG_TEXT.get("CERTIFICATES", "Zertifikate"), size=12, cell=cell
            )
        for certificate in certificates:
            start_date = certificate.get("Start Date", "")
            end_date = certificate.get("End Date", "")
            title = certificate.get("Title", "")
            institution = certificate.get("Institution", "")
            additional_info = certificate.get("Additional Information", "")
            date_text = ""
            if self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_text = f"{start_date} - {end_date}"
            elif self.cv_builder.is_value_present(start_date):
                date_text = start_date
            elif self.cv_builder.is_value_present(end_date):
                date_text = end_date
            if self.cv_builder.is_value_present(date_text):
                self.cv_builder.add_paragraph_small(
                    date_text, size=8, space_after=False, keep_with_next=True, cell=cell
                )
            if self.cv_builder.is_value_present(title):
                self.cv_builder.add_paragraph_to_cell(
                    cell,
                    title,
                    size=11,
                    bold=True,
                    space_after=False,
                    keep_with_next=True,
                )
            if self.cv_builder.is_value_present(institution):
                self.cv_builder.add_paragraph_small(
                    institution,
                    size=8,
                    bold=True,
                    space_after=False,
                    keep_with_next=False,
                    cell=cell,
                )
            if self.cv_builder.is_value_present(additional_info):
                self.cv_builder.add_paragraph_custom(
                    additional_info,
                    size=9,
                    italic=True,
                    space_after=False,
                    keep_with_next=False,
                    cell=cell,
                )
            self.cv_builder.add_spacer_to_cell(cell, space_after_pt=4)  # Adjusted space

    def add_to_cell(self, cell):
        """Add certificates section to the specified cell."""
        certificates = self.cv_builder.resume_data.get("Certificates", [])
        if self.cv_builder.is_value_present(certificates):
            # Add heading
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(LANG_TEXT.get("CERTIFICATES", "Zertifikate"))
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True

        for certificate in certificates:
            start_date = certificate.get("Start Date", "")
            end_date = certificate.get("End Date", "")
            title = certificate.get("Title", "")
            institution = certificate.get("Institution", "")
            additional_info = certificate.get("Additional Information", "")
            date_text = ""

            if self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_text = f"{start_date} - {end_date}"
            elif self.cv_builder.is_value_present(start_date):
                date_text = start_date
            elif self.cv_builder.is_value_present(end_date):
                date_text = end_date

            if self.cv_builder.is_value_present(date_text):
                # Add date text
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(date_text)
                run.font.name = "Open Sans"
                run.font.size = Pt(8)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_with_next = True

            if self.cv_builder.is_value_present(title):
                # Add title
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(title)
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(11)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_with_next = True

            if self.cv_builder.is_value_present(institution):
                # Add institution
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(institution)
                run.font.name = "Open Sans"
                run.bold = True
                run.font.size = Pt(8)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_with_next = False

            if self.cv_builder.is_value_present(additional_info):
                # Add additional information
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(additional_info)
                run.font.name = "Open Sans"
                run.italic = True
                run.font.size = Pt(9)
                paragraph.paragraph_format.space_after = Pt(4)

            # Add spacer with the same space as between languages
            self.cv_builder.add_spacer_to_cell(cell, space_after_pt=4)


class HobbiesSection:
    """Class for adding hobbies section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        """Add hobbies section to the document."""
        hobbies = self.cv_builder.resume_data.get("Hobbies", [])
        if self.cv_builder.is_value_present(hobbies):
            self.cv_builder.add_heading(LANG_TEXT.get("HOBBIES", "Hobbys"), size=12)
            hobbies_text = ", ".join(hobbies)
            self.cv_builder.add_paragraph(hobbies_text, size=11, keep_with_next=False)

    def add_to_cell(self, cell):
        """Add hobbies section to the specified cell."""
        hobbies = self.cv_builder.resume_data.get("Hobbies", [])
        if self.cv_builder.is_value_present(hobbies):
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(LANG_TEXT.get("HOBBIES", "Hobbys"))
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            hobbies_text = ", ".join(hobbies)
            paragraph = cell.add_paragraph(hobbies_text)
            for run in paragraph.runs:
                run.font.name = "Open Sans"
                run.font.size = Pt(11)
            # self.cv_builder.add_spacer_to_cell(cell, space_after_pt=10)


class SkillsAndCompetenciesSection:
    """Class for adding skills and competencies section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        """Add skills and competencies section to the document."""
        skills = self.cv_builder.resume_data.get("Skills and Competencies", {}).get(
            "Skills", []
        )
        if self.cv_builder.is_value_present(skills):
            self.cv_builder.add_heading(
                LANG_TEXT.get("SKILLS_AND_COMPETENCIES", "Fähigkeiten und Kompetenzen"),
                size=12,
            )
            skills_text = ", ".join(skills)
            self.cv_builder.add_paragraph(skills_text, size=11, keep_with_next=True)

    def add_to_cell(self, cell):
        """Add skills and competencies section to the specified cell."""
        skills = self.cv_builder.resume_data.get("Skills and Competencies", {}).get(
            "Skills", []
        )
        if self.cv_builder.is_value_present(skills):
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(
                LANG_TEXT.get("SKILLS_AND_COMPETENCIES", "Fähigkeiten und Kompetenzen")
            )
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            skills_text = ", ".join(skills)
            paragraph = cell.add_paragraph(skills_text)
            for run in paragraph.runs:
                run.font.name = "Open Sans"
                run.font.size = Pt(11)
        self.cv_builder.add_spacer_to_cell(cell, space_after_pt=10)


class SoftwareAndTechnologiesSection:
    """Class for adding software and technologies section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        """Add software and technologies section to the document."""
        technologies = self.cv_builder.resume_data.get("Software and Technologies", [])
        programming_languages = self.cv_builder.resume_data.get(
            "Skills and Competencies", {}
        ).get("Programming Languages", [])
        if self.cv_builder.is_value_present(technologies):
            self.cv_builder.add_heading(
                LANG_TEXT.get("SOFTWARE_AND_TECHNOLOGIES", "Software und Technologien"),
                size=12,
            )
            technologies_text = ", ".join(technologies)
            self.cv_builder.add_paragraph(
                technologies_text, size=11, keep_with_next=False
            )

        if self.cv_builder.is_value_present(programming_languages):
            self.cv_builder.add_heading(
                LANG_TEXT.get("PROGRAMMING_LANGUAGES", "Programmiersprachen"), size=12
            )
            for language in programming_languages:
                name = language.get("Name")
                proficiency = language.get("Proficiency Level")
                self.cv_builder.add_paragraph_custom(
                    f"{name}: {proficiency}",
                    size=11,
                    bold=False,
                    italic=False,
                    underline=False,
                    space_after=False,
                    keep_with_next=True,
                    paragraph_space_after=2,
                )

    def add_to_cell(self, cell):
        """Add software and technologies section to the specified cell."""
        technologies = self.cv_builder.resume_data.get("Software and Technologies", [])
        programming_languages = self.cv_builder.resume_data.get(
            "Skills and Competencies", {}
        ).get("Programming Languages", [])

        if self.cv_builder.is_value_present(technologies):
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(
                LANG_TEXT.get("SOFTWARE_AND_TECHNOLOGIES", "Software und Technologien")
            )
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            technologies_text = ", ".join(technologies)
            paragraph = cell.add_paragraph(technologies_text)
            for run in paragraph.runs:
                run.font.name = "Open Sans"
                run.font.size = Pt(11)
            self.cv_builder.add_spacer_to_cell(cell, space_after_pt=10)

        if self.cv_builder.is_value_present(programming_languages):
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(
                LANG_TEXT.get("PROGRAMMING_LANGUAGES", "Programmiersprachen")
            )
            run.font.name = "Open Sans"
            run.bold = True
            run.font.size = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            for language in programming_languages:
                name = language.get("Name")
                proficiency = language.get("Proficiency Level")
                paragraph = cell.add_paragraph()
                run_name = paragraph.add_run(f"{name}: ")
                run_name.bold = True
                run_name.font.name = "Open Sans"
                run_name.font.size = Pt(11)
                run_proficiency = paragraph.add_run(proficiency)
                run_proficiency.bold = False
                run_proficiency.font.name = "Open Sans"
                run_proficiency.font.size = Pt(11)
                paragraph.paragraph_format.space_after = Pt(2)
            self.cv_builder.add_spacer_to_cell(cell, space_after_pt=10)


class AdditionalInformationItemSection:
    """Class for adding an item in the additional information section."""

    def __init__(self, cv_builder, item):
        self.cv_builder = cv_builder
        self.item = item

    def add(self):
        """Add additional information item section to the document."""
        self.add_to_cell(self.cv_builder.doc.add_paragraph().add_run().element)

    def add_to_cell(self, cell):
        """Add additional information item section to the specified cell."""
        title_key = [key for key in self.item.keys() if "Title" in key]
        if title_key:
            title = self.item.get(title_key[0], "Additional Information")
            self.cv_builder.add_heading(title, size=13, cell=cell)
        for key, value in self.item.items():
            if key not in title_key and self.cv_builder.is_value_present(value):
                self.cv_builder.add_paragraph_to_cell(
                    cell, f"{key}: {value}", size=11, keep_with_next=False
                )
                self.cv_builder.add_spacer_small(cell=cell)


class AdditionalInformationSection:
    """Class for adding additional information section to the document."""

    def __init__(self, cv_builder):
        self.cv_builder = cv_builder
        self.additional_info_headline_printed = False

    def add(self):
        """Add additional information section to the document."""
        additional_info_items = self.cv_builder.resume_data.get(
            "Additional Information", []
        )
        titled_entries, untitled_entries = self._separate_entries_based_on_title(
            additional_info_items
        )
        for item in titled_entries:
            self.add_additional_info_item(item)
        if untitled_entries and not self.additional_info_headline_printed:
            self.cv_builder.add_heading(
                LANG_TEXT.get("ADDITIONAL_INFORMATION", "Zusätzliche Informationen"),
                size=13,
            )
            self.additional_info_headline_printed = True
        for item in untitled_entries:
            self.add_additional_info_item(item, include_heading=False)

    def add_additional_info_item(self, item, include_heading=True, cell=None):
        """Add an additional information item to the specified cell."""
        title = item.get("Title", "").strip()
        if self.cv_builder.is_value_present(title) and include_heading:
            self.cv_builder.add_heading(title, size=13, cell=cell)
        institution_str = item.get("Institution")
        if institution_str is not None:
            institution_str = institution_str.strip()
        else:
            institution_str = ""
        if self.cv_builder.is_value_present(institution_str):
            self.cv_builder.add_paragraph_to_cell(
                cell,
                institution_str,
                size=11,
                bold=True,
                space_after=False,
                keep_with_next=True,
            )
        date_location_str = self.format_institution_date_location(item)
        if self.cv_builder.is_value_present(date_location_str):
            self.cv_builder.add_paragraph_small(
                date_location_str,
                bold=True,
                space_after=True,
                keep_with_next=True,
                cell=cell,
            )
        address_website_str = self.format_address_website(item)
        if self.cv_builder.is_value_present(address_website_str):
            self.cv_builder.add_paragraph_custom(
                address_website_str,
                size=8,
                italic=True,
                bold=False,
                space_after=True,
                keep_with_next=True,
                cell=cell,
            )
        description = item.get("Description")
        if self.cv_builder.is_value_present(description):
            self.cv_builder.add_paragraph_to_cell(
                cell, description, size=10, space_after=True, keep_with_next=False
            )
        self.add_additional_details(item, cell=cell)
        self.cv_builder.add_spacer_large(12, cell=cell)

    def add_additional_details(self, item, cell):
        """Add additional details for an item to the specified cell."""
        for key, value in item.items():
            if key not in [
                "Title",
                "Start Date",
                "End Date",
                "Description",
                "Location",
                "Initializer",
                "Institution",
                "Address",
                "Website",
            ] and self.cv_builder.is_value_present(value):
                self.cv_builder.add_paragraph_to_cell(
                    cell, f"{key}: {value}", size=10, space_after=True
                )

    def _separate_entries_based_on_title(self, items):
        """Separate entries based on the presence of a title."""
        titled_entries = [
            item
            for item in items
            if self.cv_builder.is_value_present(item.get("Title", "").strip())
            and "Initializer" not in item
        ]
        untitled_entries = [
            item
            for item in items
            if not self.cv_builder.is_value_present(item.get("Title", "").strip())
            and "Initializer" not in item
        ]
        return titled_entries, untitled_entries

    def format_institution_date_location(self, item):
        """Format the institution, date, and location of an item."""
        institution = ""
        start_date = item.get("Start Date")
        end_date = item.get("End Date", "heute")
        location = item.get("Location")
        parts = [
            institution,
            (
                f"{start_date} - {end_date}"
                if self.cv_builder.is_value_present(start_date)
                else ""
            ),
            location,
        ]
        formatted_parts = [
            part for part in parts if self.cv_builder.is_value_present(part)
        ]
        return " | ".join(formatted_parts)

    def format_address_website(self, item):
        """Format the address and website of an item."""
        address = item.get("Address", "")
        website = item.get("Website", "")
        parts = [address, website]
        formatted_parts = [
            part for part in parts if self.cv_builder.is_value_present(part)
        ]
        return " | ".join(formatted_parts)


class Block:
    """Base class for rendering blocks in the document."""

    def render(self, cv_builder):
        pass


class HeaderBlock(Block):
    """Class for rendering a header block in the document."""

    def __init__(self, text):
        self.text = text

    def render(self, cv_builder):
        cv_builder.add_heading(self.text, size=12)


class SubheaderBlock(Block):
    """Class for rendering a subheader block in the document."""

    def __init__(self, text):
        self.text = text

    def render(self, cv_builder):
        cv_builder.add_paragraph(self.text, size=11, bold=True)


class TextBlock(Block):
    """Class for rendering a text block in the document."""

    def __init__(self, text, bold=False):
        self.text = text
        self.bold = bold

    def render(self, cv_builder):
        cv_builder.add_paragraph(self.text, size=11, bold=self.bold)


class ListBlock(Block):
    """Class for rendering a list block in the document."""

    def __init__(self, items):
        self.items = items

    def render(self, cv_builder):
        for item in self.items:
            cv_builder.add_paragraph(f"- {item}", size=11)


class DateLocationBlock(Block):
    """Class for rendering a date and location block in the document."""

    def __init__(self, date, location):
        self.date = date
        self.location = location

    def render(self, cv_builder):
        cv_builder.add_paragraph(f"{self.date} - {self.location}", size=11, bold=False)
