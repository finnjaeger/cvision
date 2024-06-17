import json
import boto3
import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, UnidentifiedImageError

# from pathlib import Path

# Import additional necessary classes or methods
debugMode = False


def lambda_handler(event, context):
    # Extract resume data and file paths from the event object
    profile_picture_present = False
    resume_data = event["resume_data"]
    s3_bucket = "cvision-completed-resumes"
    s3_key = f'{event["upload_id"]}.docx'
    logo_path = "./Branding/Sirato_Logo_Color.png"
    image_path = "/tmp/profile.png"
    footer_text = (
        "Alle auf dem Dokument enthaltenen Informationen unterliegen den Allgemeinen Geschäftsbedingungen der "
        "Sirato Recruitment GmbH.\n\n"
        "Sirato Recruitment GmbH, Geschäftsführer: René Troche | Alter Hof 5, 80331 München | Deutschland\n"
        "Registergericht: Amtsgericht München | Registernummer: HRB 276100 | Ust-ID-Nr (gemäß § 27 a "
        "Umsatzsteuergesetz): DE353433498"
    )
    contact_details_present = event.get("contact_details_present", True)

    # Download image and logo if stored in S3
    if not debugMode:
        s3 = boto3.client("s3")
        with open("/tmp/profile.png", "wb") as file:
            try:
                s3.download_fileobj(
                    "cv-profile-pictures", f'{event["upload_id"]}.png', file
                )
                profile_picture_present = True
            except Exception as e:
                print("No profile picture found")
                profile_picture_present = False
    # Example: s3.download_file('your-s3-bucket-for-input', 'logo.jpg', logo_path)
    # Example: s3.download_file('your-s3-bucket-for-input', 'profile_picture.jpg', image_path)

    # Initialize CVBuilder with the loaded data
    cv = CVBuilder(resume_data, contact_details_present)
    BorderSection(cv, logo_path, footer_text)  # Header und Footer
    # PersonalImageSection(cv, image_path).add()  # Bild
    cv.add_personal_information_with_image(
        image_path
    )  # Fügt persönliche Informationen mit Bild hinzu
    cv.add_personal_information_section(PersonalInformationSection)  # Persönliche Daten
    cv.add_section(ProfessionalSummarySection)  # Text mit einer kurzen Zusammenfassung
    cv.add_section(WorkExperienceSection)  # Berufserfahrung
    cv.add_section(EducationSection)  # Bildung
    cv.add_section(CertificatesSection)  # Zertifikate
    cv.add_section(SoftwareAndTechnologiesSection)  # Software und Technologien
    cv.add_section(SkillsAndCompetenciesSection)  # Fähigkeiten und Kompetenzen
    cv.add_section(LanguageSection)  # Sprachkenntnisse
    cv.add_section(AdditionalInformationSection)  # Zusätzliche Informationen

    # Save the document to a temporary path
    docx_path = "/tmp/resume.docx"
    cv.save_document(docx_path)

    if debugMode:
        print("DEBUG_MODE: ON")

    presigned_url = "Dummy"
    # Upload the document to S3
    if not debugMode:
        s3.upload_file(docx_path, s3_bucket, s3_key)
        # Generate a presigned URL to download the resume
        presigned_url = s3.generate_presigned_url(
            "get_object", Params={"Bucket": s3_bucket, "Key": s3_key}, ExpiresIn=3600
        )
        print(presigned_url)

    return {
        "statusCode": 200,
        "body": json.dumps({"resume": presigned_url}),
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Credentials": True,
            "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type,X-Amz-Date,Authorization,X-Api-Key,X-Amz-Security-Token",
        },
    }


# Klassendefinitionen


def set_table_cell_border(cell, **kwargs):
    """
    Sets or removes borders of a cell in a Word table.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create a 'tcBorders' element that defines the cell borders
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        side_tag = OxmlElement(f"w:{side}")
        side_tag.set(qn("w:val"), kwargs.get(side, "nil"))
        tcBorders.append(side_tag)

    tcPr.append(tcBorders)


# Spracheinstellungen

LANGUAGE_FILE = "de_DE"


def load_translation(lang_code):
    try:
        with open(f"Languages/{lang_code}.json", "r", encoding="utf-8") as file:
            print(f"Successfully loaded language file: {lang_code}.json")
            return json.load(file)
    except FileNotFoundError:
        print(
            f"Translation file not found: No such file or directory: {lang_code}.json"
        )
        return {}


LANG_TEXT = load_translation(LANGUAGE_FILE)


def convert_to_compatible_jpeg(image_path):
    output_path = "/tmp/_temp_picture.jpg"
    print("Image Path: ", image_path)
    try:
        with Image.open(image_path) as img:
            img = img.convert("RGB")
            img.save(output_path, "JPEG", quality=85, optimize=True, progressive=True)
        return True, output_path
    except Exception as e:
        return False, f"Error converting image: {e}"


def validate_image(image_path):
    try:
        with Image.open(image_path) as img:
            img.verify()
        return True, None
    except Exception as e:
        return False, f"Image validation failed for {image_path}: {e}"


# Basisklasse CVBuilder definiert allgemeine Methoden für das Dokument
class CVBuilder:

    def __init__(self, resume_data, contact_details_present=True):
        self.doc = Document()
        self.resume_data = resume_data
        self.contact_details_present = (
            contact_details_present  # Store the state of contact details visibility
        )

    def add_heading(self, text, size=15):
        heading = self.doc.add_heading(level=1)
        run = heading.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x17, 0xC1, 0x61)  # Grünton

    def add_paragraph(
        self,
        text,
        size=11,
        bold=False,
        paragraph_space_after=6,
        space_after=True,
        space_after_pt=None,
        keep_with_next=False,
    ):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold

        if space_after:
            paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
        else:
            paragraph.paragraph_format.space_after = Pt(0)
        if space_after_pt is not None:
            paragraph.paragraph_format.space_after = Pt(space_after_pt)
        paragraph.paragraph_format.keep_with_next = keep_with_next  # Hinzugefügt

    def add_paragraph_small(
        self,
        text,
        size=8,
        bold=False,
        paragraph_space_after=2,
        space_after=True,
        keep_with_next=False,
    ):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(0x69, 0x69, 0x69)  # Dunkelgrauer Ton
        if space_after:
            paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
        else:
            paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = (
            keep_with_next  # Sorgt für zusammenhängende Blöcke.
        )

    def add_paragraph_custom(
        self,
        text,
        size=8,
        bold=False,
        italic=False,
        underline=False,
        color=None,
        paragraph_space_after=2,
        space_after=True,
        keep_with_next=False,
    ):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.underline = underline

        if color:
            # Erwartet, dass die Farbe als Hex-String (z.B. "FF0000" für Rot) übergeben wird
            run.font.color.rgb = RGBColor(
                int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
            )

        if space_after:
            paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
        else:
            paragraph.paragraph_format.space_after = Pt(0)

        paragraph.paragraph_format.keep_with_next = (
            keep_with_next  # Sorgt für zusammenhängende Blöcke.
        )

    def add_paragraph_to_cell(
        self,
        cell,
        text,
        size=11,
        bold=False,
        space_after=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    ):
        """Fügt einen formatierten Absatz in eine spezifische Zelle ein."""
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Open Sans"
        run.font.size = Pt(size)
        run.bold = bold
        paragraph.alignment = (
            alignment  # Setzt die Ausrichtung basierend auf dem übergebenen Argument
        )
        if space_after:
            paragraph.paragraph_format.space_after = Pt(6)

    def add_bullet_points(
        self,
        bullet_points,
        size=10,
        bold=False,
        paragraph_space_after=2,
        space_after=True,
        keep_with_next=False,
    ):
        """
        :param bullet_points: Eine Liste von Strings, die die Bullet Points darstellen.
        :param size: Die Größe des Textes des Bullet Points. Standardmäßig 10.
        :param bold: Setzt den Text fett, wenn True. Standardmäßig False.
        :param paragraph_space_after: Der Abstand nach dem Absatz in Pt. Standardmäßig 2.
        :param space_after: Gibt an, ob nach dem Absatz ein Abstand eingefügt werden soll. Standardmäßig True.
        :param keep_with_next: Hält den Absatz mit dem nächsten zusammen, wenn True. Standardmäßig False.
        """
        for bullet_point in bullet_points:
            paragraph = self.doc.add_paragraph(style="ListBullet")
            run = paragraph.add_run(bullet_point)
            run.font.name = "Open Sans"
            run.font.size = Pt(size)
            run.bold = bold
            if space_after:
                paragraph.paragraph_format.space_after = Pt(paragraph_space_after)
            else:
                paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = keep_with_next

    def add_spacer_small(self, space_after_pt=2):
        if self.doc.paragraphs:
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.paragraph_format.space_after = Pt(space_after_pt)

    def add_spacer_medium(self, space_after_pt=4):
        if self.doc.paragraphs:
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.paragraph_format.space_after = Pt(space_after_pt)

    def add_spacer_large(self, space_after_pt=6):
        if self.doc.paragraphs:
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.paragraph_format.space_after = Pt(space_after_pt)

    def save_document(self, path):
        self.doc.save(path)
        print(f"Successfully created the document: {path}")

    # Zentralisierte Methode zum Hinzufügen von Abschnitten
    def add_section(self, section_class, *args, **kwargs):
        # Instantiate the section class with additional arguments
        section = section_class(self, *args, **kwargs)
        section.add()

    def is_value_present(self, value):
        """Prüft, ob ein Wert aussagekräftig ist und nicht 'NULL', 'N/A' oder leer."""
        if value in [None, "null", "NULL", "N/A", "", []]:
            return False
        return True

    def add_personal_information_section(self, section_class):
        section = section_class(self, self.contact_details_present)
        section.add()

    def add_personal_information(self, cell):
        """Adds personal information directly into a specified table cell."""
        personal_info = self.resume_data.get("Personal Information", {})

        # Full name
        full_name = f"{personal_info.get('Firstname', '')} {personal_info.get('Surname', '')}".strip()
        if full_name:
            self.add_paragraph_to_cell(
                cell,
                full_name,
                size=17,
                bold=True,
                space_after=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )

        # Current role
        if self.is_value_present(personal_info.get("Current Role")):
            self.add_paragraph_to_cell(
                cell,
                f"{personal_info.get('Current Role')}",
                size=10,
                bold=False,
                space_after=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )

    def add_personal_information_with_image(self, image_path):
        # Erstellt eine Tabelle mit einer Reihe und drei Spalten
        table = self.doc.add_table(rows=1, cols=3)
        table.columns[0].width = Inches(
            4.0
        )  # Breite für die Spalte mit persönlichen Informationen anpassen
        table.columns[1].width = Inches(
            0.5
        )  # Mittlere Spalte als Abstandshalter, möglichst klein halten
        table.columns[2].width = Inches(
            2.0
        )  # Breite für die Spalte mit dem Bild anpassen

        # Stellt die Zelleneinstellungen ein
        for cell in table.columns[0].cells:
            cell.width = Inches(4.4)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.1)
        for cell in table.columns[2].cells:
            cell.width = Inches(2.0)

        # Formatierung der Zellen zur horizontalen und vertikalen Ausrichtung
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in table.rows[0].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Einfügen der persönlichen Informationen in die erste Zelle (linke Spalte)
        self.add_personal_information(table.cell(0, 0))

        # Einfügen des Bildes in die dritte Zelle (rechte Spalte), falls ein Bildpfad vorhanden ist
        if self.is_value_present(image_path):
            cell = table.cell(0, 2)
            paragraph = cell.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)

            success, result_or_error = convert_to_compatible_jpeg(image_path)
            if success:
                valid, error = validate_image(result_or_error)
                if valid:
                    run.add_picture(result_or_error, width=Inches(1.5))
                else:
                    print(error)  # Handle image validation error
            else:
                print(result_or_error)  # Handle image conversion error

        # Entfernen der Tabellenränder
        for cell in table.rows[0].cells:
            set_table_cell_border(
                cell,
                top="nil",
                left="nil",
                bottom="nil",
                right="nil",
                insideH="nil",
                insideV="nil",
            )

    def add_additional_information_blocks(self, blocks):
        block_mapping = {
            "HeaderBlock": HeaderBlock,
            "SubheaderBlock": SubheaderBlock,
            "TextBlock": TextBlock,
            "ListBlock": ListBlock,
            "DateLocationBlock": DateLocationBlock,
        }

        for block in blocks:
            block_type = block.get("type")
            block_class = block_mapping.get(block_type)

            if block_class:
                # Erzeugung und Rendering des Block-Objekts
                if block_type in ["HeaderBlock", "SubheaderBlock", "TextBlock"]:
                    content = block.get("content", "")
                    bold = block.get("bold", False)
                    block_instance = (
                        block_class(content, bold=bold)
                        if block_type == "TextBlock"
                        else block_class(content)
                    )
                elif block_type == "ListBlock":
                    items = block.get("items", [])
                    block_instance = block_class(items)
                elif block_type == "DateLocationBlock":
                    date = block.get("date", "")
                    location = block.get("location", "")
                    block_instance = block_class(date, location)

                block_instance.render(self)


# Spezialisierte Klassen für jeden Abschnitt


class BorderSection:
    def __init__(
        self, cv_builder, logo_path, footer_text=None, footer_distance=Inches(0.2)
    ):
        self.cv_builder = cv_builder
        self.logo_path = logo_path
        self.footer_text = footer_text
        self.footer_distance = footer_distance
        self.add_header()
        if self.footer_text:
            self.add_footer()

    def add_header(self):
        header = self.cv_builder.doc.sections[0].header
        paragraph = (
            header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        )
        run = paragraph.add_run()
        run.add_picture(self.logo_path, width=Pt(100))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_footer(self):
        section = self.cv_builder.doc.sections[
            0
        ]  # Geht davon aus, dass Ihr Dokument nur eine Sektion hat
        footer = section.footer
        paragraph = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )

        # Fügt den Trennstrich und den Text hinzu
        run = paragraph.add_run()
        run.add_text("_" * 175)  # Erzeugt einen horizontalen Strich
        paragraph.add_run("\n" + self.footer_text)

        # Formatierung
        for run in paragraph.runs:
            run.font.size = Pt(6)  # Stellt die Schriftgröße ein

        # Zentriert den Footer-Text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Setzt den Abstand des Footers vom unteren Rand der Seite
        section.footer_distance = self.footer_distance


# class PersonalImageSection:
#     def __init__(self, cv_builder, image_path):
#         self.cv_builder = cv_builder
#         self.original_image_path = image_path

#     def add(self):
#         compatible_image_path = convert_to_compatible_jpeg(self.original_image_path)
#         if compatible_image_path and validate_image(compatible_image_path):
#             try:
#                 self.cv_builder.doc.add_picture(compatible_image_path, width=Pt(100))
#                 print("Image successfully added to the document.")
#             except Exception as e:
#                 print(f"Failed to add image to the document: {e}")
#         else:
#             print("Image conversion failed or resulted in an unsupported format.")


def is_value_present(value):
    """Prüft, ob ein Wert aussagekräftig ist und nicht 'NULL', 'N/A' oder leer."""
    if value in [None, "null", "NULL", "N/A", "", []]:
        return False
    return True


class PersonalInformationSection:
    def __init__(self, cv_builder, contact_details_present):
        self.cv_builder = cv_builder
        self.contact_details_present = contact_details_present

    def add(self):
        personal_info = self.cv_builder.resume_data["Personal Information"]
        contact_info = self.cv_builder.resume_data["Contact Information"]

        # Überprüfe, ob Informationen vorhanden sind, die unter der Überschrift angezeigt werden sollen
        info_present = any(
            [
                is_value_present(personal_info.get("Current Role")),
                is_value_present(personal_info.get("Birthday")),
                is_value_present(personal_info.get("Nationality")),
                is_value_present(personal_info.get("Marital Status")),
                is_value_present(personal_info.get("Availability")),
                is_value_present(personal_info.get("Additional Information")),
            ]
        )

        if info_present:
            # Füge die Überschrift nur hinzu, wenn Informationen vorhanden sind
            self.cv_builder.add_heading(
                LANG_TEXT.get("PERSONAL_INFORMATION", "Persönliche Informationen"),
                size=13,
            )

            if is_value_present(personal_info.get("Birthday")):
                self.cv_builder.add_paragraph(
                    f"{LANG_TEXT.get('BIRTHDAY', 'Geburtsdatum')}: {personal_info.get('Birthday')}",
                    size=11,
                )

            if is_value_present(personal_info.get("Nationality")):
                self.cv_builder.add_paragraph(
                    f"{LANG_TEXT.get('NATIONALITY', 'Nationalität')}: {personal_info.get('Nationality')}",
                    size=11,
                )

            if is_value_present(personal_info.get("Marital Status")):
                self.cv_builder.add_paragraph(
                    f"{LANG_TEXT.get('MARITAL_STATUS', 'Familienstand')}: {personal_info.get('Marital Status')}",
                    size=11,
                )

            if is_value_present(personal_info.get("Availability")):
                self.cv_builder.add_paragraph(
                    f"{LANG_TEXT.get('AVAILABILITY', 'Verfügbarkeit')}: {personal_info.get('Availability')}",
                    size=11,
                )

            if is_value_present(personal_info.get("Additional Information")):
                self.cv_builder.add_paragraph_custom(
                    f"{LANG_TEXT.get('ADDITIONAL_INFORMATION', 'Zusätzliche Informationen')}: {personal_info.get('Additional Information')}",
                    size=9,
                    italic=True,
                )

        # Contact Information
        if self.contact_details_present:
            self.cv_builder.add_heading(LANG_TEXT.get("CONTACT", "Kontakt"), size=13)

            if is_value_present(contact_info.get("Address")):
                self.cv_builder.add_paragraph(
                    f"{LANG_TEXT.get('ADDRESS', 'Adresse')}: {contact_info.get('Address')}",
                    size=11,
                )

            phone_numbers = [
                number
                for key, number in contact_info.items()
                if key.endswith("Phone Number") and is_value_present(number)
            ]
            if phone_numbers:
                self.cv_builder.add_paragraph(
                    f"{LANG_TEXT.get('PHONE', 'Telefon')}: {' | '.join(phone_numbers)}",
                    size=11,
                )

            if is_value_present(contact_info.get("Email")):
                self.cv_builder.add_paragraph(
                    f"{LANG_TEXT.get('EMAIL', 'E-Mail')}: {contact_info.get('Email')}",
                    size=11,
                )

            if is_value_present(contact_info.get("Additional Information")):
                self.cv_builder.add_paragraph_custom(
                    f"{LANG_TEXT.get('MORE_CONTACT_INFO', 'Weitere Kontaktinformationen')}: {contact_info.get('Additional Information')}",
                    size=9,
                    italic=True,
                )


class ProfessionalSummarySection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        # Zugriff auf den 'Professional Summary' Bereich im resume_data
        summary_data = self.cv_builder.resume_data.get("Professional Summary", {})
        summary_text = summary_data.get("Professional Summary Text", "")
        summary_bullet_points = summary_data.get(
            "Professional Summary Bullet Points", []
        )

        if self.cv_builder.is_value_present(summary_text):
            self.cv_builder.add_heading(
                LANG_TEXT.get("PROFESSIONAL_SUMMARY", "Berufliche Zusammenfassung"),
                size=13,
            )
            self.cv_builder.add_paragraph(
                summary_text, size=11, keep_with_next=True
            )  # Stellt sicher, dass der Text und die Bullet Points als zusammengehöriger Block behandelt werden

        if summary_bullet_points:
            self.cv_builder.add_bullet_points(
                summary_bullet_points, size=10, keep_with_next=False
            )  # Fügt die Bullet Points hinzu


class WorkExperienceSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        self.cv_builder.add_heading(
            LANG_TEXT.get("WORK_EXPERIENCE", "Berufserfahrung"), size=13
        )
        experiences = self.cv_builder.resume_data.get("Working Experience", [])
        for item in experiences:
            title = item.get("Title", "NULL")
            company = item.get("Company", "")
            location = item.get("Location", "")
            start_date = item.get("Start Date", "")
            end_date = item.get("End Date", "heute")
            description = item.get("Description", "")
            bullet_points = item.get("Bullet Points", [])
            website = item.get("Website", "")
            additional_information = item.get("Additional Information", "")

            # Format Company and Location (Wird vereinfacht via "company_and_location_and_website" repräsentiert)

            company_and_location_and_website = ""
            # Überprüfen und Hinzufügen von Company, wenn vorhanden
            if self.cv_builder.is_value_present(company):
                company_and_location_and_website += item.get("Company", "")
            # Überprüfen und Hinzufügen von Location, wenn vorhanden und Company bereits hinzugefügt wurde
            if self.cv_builder.is_value_present(location):
                if (
                    company_and_location_and_website
                ):  # Wenn bereits ein Unternehmensname vorhanden ist
                    company_and_location_and_website += " | "
                company_and_location_and_website += item.get("Location", "")
            # Überprüfen und Hinzufügen von Website, wenn vorhanden und einer der anderen Werte bereits hinzugefügt
            # wurde
            if self.cv_builder.is_value_present(website):
                if (
                    company_and_location_and_website
                ):  # Wenn bereits ein Unternehmensname oder Standort vorhanden ist
                    company_and_location_and_website += " | "
                company_and_location_and_website += item.get("Website", "")

            # Format Date (Wird vereinfacht via "date_range" präsentiert)
            date_range = start_date
            if self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range += f" - {end_date}"
            elif not self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range = end_date

            if self.cv_builder.is_value_present(date_range):
                self.cv_builder.add_paragraph_small(
                    date_range, size=8, space_after=False, keep_with_next=True
                )
            if self.cv_builder.is_value_present(title):
                self.cv_builder.add_paragraph(
                    title, size=11, bold=True, space_after=False, keep_with_next=True
                )
            if self.cv_builder.is_value_present(company_and_location_and_website):
                self.cv_builder.add_paragraph_small(
                    company_and_location_and_website,
                    space_after=False,
                    bold=True,
                    keep_with_next=True,
                )
            if self.cv_builder.is_value_present(description):
                self.cv_builder.add_paragraph(
                    description, size=10, space_after=False, keep_with_next=True
                )

            if bullet_points:
                self.cv_builder.add_spacer_small()
                self.cv_builder.add_bullet_points(bullet_points, keep_with_next=False)

            if self.cv_builder.is_value_present(additional_information):
                self.cv_builder.add_paragraph_custom(
                    additional_information,
                    size=9,
                    italic=True,
                    space_after=False,
                    keep_with_next=False,
                )
            self.cv_builder.add_spacer_large(space_after_pt=12)


class EducationSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        self.cv_builder.add_heading(LANG_TEXT.get("EDUCATION", "Bildung"), size=13)
        education_items = self.cv_builder.resume_data.get("Education", [])
        for item in education_items:
            diploma = item.get("Diploma", "NULL")
            institution = item.get("Institution", "")
            start_date = item.get("Start Date", "")
            end_date = item.get("End Date", "heute")
            grade = item.get("Grade", "")
            location = item.get("Location", "")
            website = item.get("Website", "")
            description = item.get("Description", "")
            bullet_points = item.get("Bullet Points", [])
            additional_information = item.get("Additional Information")

            institution_and_location_and_website = ""
            # Überprüfen und Hinzufügen von Company, wenn vorhanden
            if self.cv_builder.is_value_present(institution):
                institution_and_location_and_website += item.get("Institution", "")
            # Überprüfen und Hinzufügen von Location, wenn vorhanden und Company bereits hinzugefügt wurde
            if self.cv_builder.is_value_present(location):
                if (
                    institution_and_location_and_website
                ):  # Wenn bereits ein Unternehmensname vorhanden ist
                    institution_and_location_and_website += " | "
                institution_and_location_and_website += item.get("Location", "")
            # Überprüfen und Hinzufügen von Website, wenn vorhanden und einer der anderen Werte bereits hinzugefügt
            # wurde
            if self.cv_builder.is_value_present(website):
                if (
                    institution_and_location_and_website
                ):  # Wenn bereits ein Unternehmensname oder Standort vorhanden ist
                    institution_and_location_and_website += " | "
                institution_and_location_and_website += item.get("Website", "")

            date_range = start_date
            if self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range += f" - {end_date}"
            elif not self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_range = end_date

            if self.cv_builder.is_value_present(date_range):
                self.cv_builder.add_paragraph_small(
                    date_range, size=8, space_after=False, keep_with_next=True
                )
            if self.cv_builder.is_value_present(diploma):
                self.cv_builder.add_paragraph(
                    diploma, size=11, bold=True, space_after=False, keep_with_next=True
                )
            if self.cv_builder.is_value_present(institution_and_location_and_website):
                self.cv_builder.add_paragraph_small(
                    institution_and_location_and_website,
                    space_after=False,
                    bold=True,
                    keep_with_next=True,
                )
            if self.cv_builder.is_value_present(grade):
                self.cv_builder.add_paragraph(
                    LANG_TEXT.get("GRADE", "Note") + " " + grade,
                    size=10,
                    space_after=False,
                    keep_with_next=False,
                )

            if self.cv_builder.is_value_present(description):
                self.cv_builder.add_paragraph(
                    description, size=10, space_after=False, keep_with_next=True
                )

            if bullet_points:
                self.cv_builder.add_spacer_small()
                self.cv_builder.add_bullet_points(bullet_points, keep_with_next=True)

            if self.cv_builder.is_value_present(additional_information):
                self.cv_builder.add_paragraph_custom(
                    additional_information,
                    size=9,
                    italic=True,
                    space_after=False,
                    keep_with_next=False,
                )
            self.cv_builder.add_spacer_large(space_after_pt=12)


class LanguageSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        languages = self.cv_builder.resume_data.get("Languages", [])
        if languages:
            self.cv_builder.add_heading(LANG_TEXT.get("LANGUAGES", "Sprachen"), size=13)
            for language in languages:
                self.add_language_paragraph(language["Name"], language["Level"])

    def add_language_paragraph(self, name, level):
        paragraph = self.cv_builder.doc.add_paragraph()
        run = paragraph.add_run(f"{name}: ")
        run.font.name = "Open Sans"  # für gesprochene Sprache: z.B. '"Englisch": XXXXX'
        run.bold = True
        run.font.size = Pt(11)
        run = paragraph.add_run(
            level
        )  # für das eigentliche Sprachlevel: z.B. 'XXXXX: "Muttersprachler"'
        run.font.name = "Open Sans"
        run.bold = False
        run.font.size = Pt(11)
        paragraph.paragraph_format.space_after = Pt(4)


class CertificatesSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):

        certificates = self.cv_builder.resume_data.get("Certificates", [])

        if certificates:  # Wenn `certificates` nicht leer ist
            self.cv_builder.add_heading(
                LANG_TEXT.get("CERTIFICATES", "Zertifikate"), size=13
            )

        for certificate in certificates:
            start_date = certificate.get("Start Date", "")
            end_date = certificate.get("End Date", "")
            title = certificate.get("Title", "")
            institution = certificate.get("Institution", "")

            date_text = ""
            if self.cv_builder.is_value_present(
                start_date
            ) and self.cv_builder.is_value_present(end_date):
                date_text = f"{start_date} - {end_date}"
            elif self.cv_builder.is_value_present(start_date):
                date_text = start_date
            elif self.cv_builder.is_value_present(
                end_date
            ):  # Wenn nur End Date vorhanden ist, wird das verwendet
                date_text = end_date

            if self.cv_builder.is_value_present(date_text):
                self.cv_builder.add_paragraph_small(
                    date_text, size=8, space_after=False, keep_with_next=True
                )
            if self.cv_builder.is_value_present(title):
                self.cv_builder.add_paragraph(
                    title, size=11, bold=True, space_after=False, keep_with_next=True
                )
            if self.cv_builder.is_value_present(institution):
                self.cv_builder.add_paragraph_small(
                    institution, bold=True, space_after=False, keep_with_next=False
                )
            self.cv_builder.add_spacer_large()


class HobbiesSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        hobbies = self.cv_builder.resume_data.get("Hobbies", [])
        if self.cv_builder.is_value_present(hobbies):
            self.cv_builder.add_heading(LANG_TEXT.get("HOBBIES", "Hobbys"), size=13)
            hobbies_text = ", ".join(hobbies)
            self.cv_builder.add_paragraph(hobbies_text, size=11, keep_with_next=False)


class SkillsAndCompetenciesSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        skills = self.cv_builder.resume_data.get("Skills and Competencies", [])
        if self.cv_builder.is_value_present(skills):
            self.cv_builder.add_heading(
                LANG_TEXT.get("SKILLS_AND_COMPETENCIES", "Fähigkeiten und Kompetenzen"),
                size=13,
            )
            skills_text = "\n".join(skills)
            self.cv_builder.add_paragraph(skills_text, size=11, keep_with_next=False)


class SoftwareAndTechnologiesSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder

    def add(self):
        technologies = self.cv_builder.resume_data.get("Software and Technologies", [])
        if self.cv_builder.is_value_present(technologies):
            self.cv_builder.add_heading(
                LANG_TEXT.get("SOFTWARE_AND_TECHNOLOGIES", "Software und Technologien"),
                size=13,
            )
            technologies_text = ", ".join(technologies)
            self.cv_builder.add_paragraph(
                technologies_text, size=11, keep_with_next=False
            )


class AdditionalInformationItemSection:
    def __init__(self, cv_builder, item):
        self.cv_builder = cv_builder
        self.item = item

    def add(self):
        # Verwendung des Titels als Überschrift, wenn vorhanden
        title_key = [key for key in self.item.keys() if "Title" in key]
        if title_key:
            title = self.item.get(title_key[0], "Additional Information")
            self.cv_builder.add_heading(title, size=13)

        # Iteration über alle Schlüssel-Wert-Paare im Item, außer dem Titel
        for key, value in self.item.items():
            if key not in title_key and self.cv_builder.is_value_present(value):
                # Hier könnten Sie entscheiden, wie Sie mit bestimmten Schlüsseln umgehen wollen
                # Zum Beispiel könnte man 'Description' anders formatieren als 'Start Date'
                self.cv_builder.add_paragraph(
                    f"{key}: {value}", size=11, keep_with_next=False
                )
                self.cv_builder.add_spacer_small()


class AdditionalInformationSection:
    def __init__(self, cv_builder):
        self.cv_builder = cv_builder
        self.additional_info_headline_printed = False

    def add(self):
        # Zuerst die Einträge aus dem Resume-Daten extrahieren
        additional_info_items = self.cv_builder.resume_data.get(
            "Additional Information", []
        )

        # Vorsortierung: Trennen von Einträgen mit und ohne Titel
        titled_entries, untitled_entries = self._separate_entries_based_on_title(
            additional_info_items
        )

        # Zuerst Einträge mit Titeln verarbeiten
        for item in titled_entries:
            self.add_additional_info_item(item)

        # Dann "Weitere Informationen" als einmalige Überschrift, wenn es untitled Einträge gibt
        if untitled_entries and not self.additional_info_headline_printed:
            self.cv_builder.add_heading(
                LANG_TEXT.get("ADDITIONAL_INFORMATION", "Zusätzliche Informationen"),
                size=13,
            )
            self.additional_info_headline_printed = True

        # Verarbeitung der untitled Einträge
        for item in untitled_entries:
            self.add_additional_info_item(item, include_heading=False)

    def add_additional_info_item(self, item, include_heading=True):
        # Überprüft, ob ein Titel vorhanden ist, und entfernt eventuelle Leerzeichen.
        title = item.get("Title", "").strip()

        # Wenn ein Titel vorhanden ist und include_heading True ist, diesen verwenden.
        if self.cv_builder.is_value_present(title) and include_heading:
            self.cv_builder.add_heading(title, size=13)

        # Wenn eine Institution vorhanden ist, wird diese als "Unterüberschrift" verwendet.
        institution_str = item.get("Institution")
        if institution_str is not None:
            institution_str = institution_str.strip()
        else:
            institution_str = ""

        if self.cv_builder.is_value_present(institution_str):
            self.cv_builder.add_paragraph(
                institution_str,
                size=11,
                bold=True,
                space_after=False,
                keep_with_next=True,
            )

        # Datum und Ort formatieren, wenn vorhanden
        date_location_str = self.format_institution_date_location(item)
        if self.cv_builder.is_value_present(date_location_str):
            self.cv_builder.add_paragraph_small(
                date_location_str, bold=True, space_after=True, keep_with_next=True
            )

        # Adresse und Website formatieren, wenn vorhanden
        address_website_str = self.format_address_website(item)
        if self.cv_builder.is_value_present(address_website_str):
            self.cv_builder.add_paragraph_custom(
                address_website_str,
                size=8,
                italic=True,
                bold=False,
                space_after=True,
                keep_with_next=True,
            )

        # Beschreibung, wenn vorhanden
        description = item.get("Description")
        if self.cv_builder.is_value_present(description):
            self.cv_builder.add_paragraph(
                description, size=10, space_after=True, keep_with_next=False
            )

        # Weitere Details
        self.add_additional_details(item)

        # Fügt einen Spacer hinzu, um Abschnitte zu trennen
        self.cv_builder.add_spacer_large(12)

    def _separate_entries_based_on_title(self, items):
        # Trennen der Einträge in solche mit und ohne Titel
        titled_entries = [
            item
            for item in items
            if self.cv_builder.is_value_present(item.get("Title", "").strip())
            and "Initializer" not in item
        ]
        untitled_entries = [
            item
            for item in items
            if not self.cv_builder.is_value_present(item.get("Title", "").strip())
            and "Initializer" not in item
        ]
        return titled_entries, untitled_entries

    def format_institution_date_location(self, item):
        # Extrahiert Institution, Start- und Enddatum sowie Ort
        # institution = item.get("Institution")
        institution = ""  # TEMPORARLY FOR DESIGN TESTING.
        start_date = item.get("Start Date")
        end_date = item.get("End Date", "heute")
        location = item.get("Location")

        # Formatierung des Datumsbereichs und Orts in einer Zeile
        parts = [
            institution,
            (
                f"{start_date} - {end_date}"
                if self.cv_builder.is_value_present(start_date)
                else ""
            ),
            location,
        ]
        formatted_parts = [
            part for part in parts if self.cv_builder.is_value_present(part)
        ]
        return " | ".join(formatted_parts)

    def format_address_website(self, item):
        # Extrahiert Adresse und Website und formatiert diese
        address = item.get("Address", "")
        website = item.get("Website", "")
        parts = [address, website]
        formatted_parts = [
            part for part in parts if self.cv_builder.is_value_present(part)
        ]
        return " | ".join(formatted_parts)

    def add_additional_details(self, item):
        # Verarbeitet alle zusätzlichen Detailinformationen, ausgenommen speziell behandelte Schlüssel
        for key, value in item.items():
            if key not in [
                "Title",
                "Start Date",
                "End Date",
                "Description",
                "Location",
                "Initializer",
                "Institution",
                "Address",
                "Website",
            ] and self.cv_builder.is_value_present(value):
                self.cv_builder.add_paragraph(
                    f"{key}: {value}", size=10, space_after=True
                )


# Additional Information Sektion (erweitert) START
class Block:
    def render(self, cv_builder):
        pass


class HeaderBlock(Block):
    def __init__(self, text):
        self.text = text

    def render(self, cv_builder):
        cv_builder.add_heading(self.text, size=12)


class SubheaderBlock(Block):
    def __init__(self, text):
        self.text = text

    def render(self, cv_builder):
        cv_builder.add_paragraph(self.text, size=11, bold=True)


class TextBlock(Block):
    def __init__(self, text, bold=False):
        self.text = text
        self.bold = bold

    def render(self, cv_builder):
        cv_builder.add_paragraph(self.text, size=11, bold=self.bold)


class ListBlock(Block):
    def __init__(self, items):
        self.items = items

    def render(self, cv_builder):
        for item in self.items:
            cv_builder.add_paragraph(f"- {item}", size=11)


class DateLocationBlock(Block):
    def __init__(self, date, location):
        self.date = date
        self.location = location

    def render(self, cv_builder):
        cv_builder.add_paragraph(f"{self.date} - {self.location}", size=11, bold=False)


# Additional Information Sektion (erweitert) ENDE

# Wenn weitere Bereiche im CV hinzugefügt werden (z. B. ein Abschnitt für Software-Skills, soll diese hier als neue
# Klasse implementiert werden.
